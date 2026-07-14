from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "documentation"
OUT = DOCS / "Business_Card_Scanner_Submission.docx"
PROCESS_FLOW = DOCS / "process-flow.png"
ER_DIAGRAM = DOCS / "er-diagram.png"


BLUE = RGBColor(31, 78, 121)
TEXT = RGBColor(31, 31, 31)
MUTED = RGBColor(90, 90, 90)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2EC"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_text(paragraph, text, bold=False, italic=False, color=TEXT):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    return run


def add_paragraph(doc, text="", style=None):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    if text:
        add_text(paragraph, text)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(level=level)
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.color.rgb = BLUE if level <= 2 else TEXT
    run.bold = True
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        add_text(paragraph, item)


def add_numbers(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        add_text(paragraph, item)


def add_code_block(doc, lines):
    for line in lines:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.paragraph_format.right_indent = Inches(0.2)
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(40, 40, 40)


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    set_table_borders(table)
    for label, value in rows:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_width(row.cells[0], 2500)
        set_cell_width(row.cells[1], 6860)
        set_cell_shading(row.cells[0], "F3F6FA")
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = TEXT
            row.cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()
    return table


def add_image(doc, image_path, caption):
    if not image_path.exists():
        add_paragraph(doc, "Diagram missing: {}".format(image_path.name))
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    if image_path.name == "process-flow.png":
        run.add_picture(str(image_path), height=Inches(8.25))
    else:
        run.add_picture(str(image_path), width=Inches(6.35))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    add_text(cap, caption, italic=True, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(10.5)


def build():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title_run = title.add_run("Business Card Scanner & Smart Contact Manager")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(18, 52, 86)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    add_text(
        subtitle,
        "Submission document for the NestJS, PostgreSQL, Drizzle ORM, and React backend developer assignment.",
        color=MUTED,
    )

    add_heading(doc, "Submission Links", 1)
    add_key_value_table(
        doc,
        [
            ("Recorded demo video", "PASTE_GOOGLE_DRIVE_OR_ZOHO_LINK_HERE"),
            ("Source code archive", "PASTE_GOOGLE_DRIVE_OR_ZOHO_LINK_HERE"),
        ],
    )

    add_heading(doc, "Project Summary", 1)
    add_paragraph(
        doc,
        "This project is a contact manager built around two fast entry methods: business card scanning and voice input. A user can scan a card, speak contact details, or type details manually. The app then lets the user review the contact, checks for duplicates, supports merge or create-new decisions, stores relationships and groups, exports VCF files, and shows a relationship graph.",
    )
    add_paragraph(
        doc,
        "The implementation keeps OCR and speech processing local. The goal is not to hide uncertainty from the user; the app fills a draft first and the user confirms the result before it becomes a saved contact.",
    )

    add_heading(doc, "Reviewer Quick Start", 1)
    add_paragraph(doc, "The easiest way to run the project is Docker:")
    add_code_block(doc, ["docker compose up --build"])
    add_paragraph(doc, "Open the frontend:")
    add_code_block(doc, ["http://localhost:5173"])
    add_key_value_table(
        doc,
        [
            ("Frontend", "http://localhost:5173"),
            ("Backend API", "http://localhost:3000"),
            ("Adminer", "http://localhost:8080"),
            ("Database", "PostgreSQL: bhumio / bhumio_dev_password / bhumio_contacts"),
        ],
    )
    add_paragraph(
        doc,
        "The API container runs Drizzle migrations on startup. Sample data is loaded only when the database is empty, so contacts created during review are not wiped by a normal container restart.",
    )
    add_paragraph(doc, "If the default ports are already in use:")
    add_code_block(
        doc,
        [
            "API_PORT=3002 WEB_PORT=5174 docker compose up --build",
            "",
            "# PowerShell",
            "$env:API_PORT='3002'; $env:WEB_PORT='5174'; docker compose up --build",
        ],
    )

    add_heading(doc, "Implemented Requirements", 1)
    add_heading(doc, "Business Card Scanner", 2)
    add_paragraph(doc, "Implemented with local Tesseract OCR. The card scan can populate:")
    add_bullets(
        doc,
        [
            "Full name",
            "Designation or title",
            "Company",
            "Email",
            "Phone number",
            "Website",
            "Address",
            "Business relationship",
        ],
    )

    add_heading(doc, "Voice-Based Contact Entry", 2)
    add_paragraph(
        doc,
        "Implemented using browser audio recording and a local faster-whisper runner on the backend. No cloud speech API is required. The transcript is parsed into the same contact form used by scanned cards.",
    )

    add_heading(doc, "Duplicate Detection", 2)
    add_paragraph(doc, "Duplicate checks run automatically while creating contacts and use:")
    add_bullets(doc, ["Name", "Email", "Phone"])
    add_paragraph(doc, "When a possible match is found, the UI offers Use Existing, Merge, or Create New.")

    add_heading(doc, "Relationships and Groups", 2)
    add_paragraph(
        doc,
        "The app supports direct contact-to-contact relationships, named groups, contacts belonging to multiple groups, and a graph view with relationship labels and group coloring.",
    )

    add_heading(doc, "Contact Management", 2)
    add_bullets(
        doc,
        [
            "Contact list and detail view",
            "Single and multi-contact VCF export",
            "Manual relationship linking after save",
            "Group assignment from the contact detail page",
            "Soft delete for contacts",
            "Relationship graph with group filter and search",
        ],
    )

    add_heading(doc, "Application Flow", 1)
    add_paragraph(
        doc,
        "Both input methods lead into the same review and save flow. The diagram below is the same Mermaid-rendered process diagram included in the project documentation.",
    )
    add_image(doc, PROCESS_FLOW, "Application process flow")

    add_heading(doc, "Database ER Diagram", 1)
    add_paragraph(
        doc,
        "The schema keeps contacts, contact methods, relationships, groups, and extraction history separate so the data remains flexible as the product grows.",
    )
    add_image(doc, ER_DIAGRAM, "Database ER diagram")

    add_heading(doc, "Database Design", 1)
    add_paragraph(doc, "Main tables:")
    add_bullets(
        doc,
        [
            "contacts",
            "contact_emails",
            "contact_phones",
            "contact_websites",
            "contact_addresses",
            "contact_relationships",
            "contact_groups",
            "contact_group_members",
            "extraction_attempts",
        ],
    )
    add_paragraph(
        doc,
        "One contact can have many emails, phone numbers, websites, and addresses. Contacts can link to other contacts through contact_relationships, and they can belong to many groups through contact_group_members.",
    )

    add_heading(doc, "API Design", 1)
    add_paragraph(doc, "Important REST endpoints:")
    add_bullets(
        doc,
        [
            "POST /extractions/business-card",
            "POST /extractions/voice",
            "GET /contacts",
            "POST /contacts",
            "GET /contacts/:id",
            "DELETE /contacts/:id",
            "POST /contacts/duplicates/check",
            "POST /contacts/:id/merge",
            "POST /contacts/:id/relationships",
            "GET /contacts/groups",
            "POST /contacts/:id/groups",
            "GET /contacts/graph",
            "GET /contacts/:id/vcf",
            "GET /contacts/export/vcf?ids=id1,id2",
        ],
    )

    add_heading(doc, "Manual Setup", 1)
    add_code_block(
        doc,
        [
            "npm install",
            "npm run dev:infra",
            "npm run db:migrate",
            "npm run db:seed",
            "npm run dev:api",
            "npm run dev:web -- --host 0.0.0.0",
        ],
    )
    add_key_value_table(doc, [("Frontend", "http://localhost:5173"), ("Backend", "http://localhost:3000")])

    add_heading(doc, "Tests and Verification", 1)
    add_paragraph(doc, "Useful commands:")
    add_code_block(
        doc,
        [
            "npm run test:unit",
            "npm run test:e2e",
            "npm run build",
            "npm run test:ocr",
            "npm run test",
        ],
    )
    add_paragraph(
        doc,
        "The OCR evaluation uses the labeled 20-card dataset under datasets/business-cards and prints a field-by-field CLI report.",
    )

    add_heading(doc, "Sample Data", 1)
    add_paragraph(doc, "Sample data is provided through:")
    add_code_block(doc, ["npm run db:seed"])
    add_bullets(
        doc,
        [
            "Bhumio group: Tejas Kamal Sahoo, Aarav Mehta, Vinod C",
            "Doe Family group: John Doe, Sarah Doe, Jane Doe, Jack Doe",
            "Work partner links inside the Bhumio group",
            "Father, mother, son, daughter, sister, and brother links inside the Doe Family group",
        ],
    )

    add_heading(doc, "Packaging Checklist", 1)
    add_bullets(
        doc,
        [
            "Include source code, documentation, docs, labels.csv, and sample data scripts.",
            "Include .env.example files, not private .env files.",
            "Exclude node_modules, dist, local logs, and Docker volumes.",
            "Keep the demo video link and source archive link public.",
            "Do not submit a GitHub or GitLab repository link.",
            "In the demo, show card scan, voice fill, duplicate handling, relationship linking, graph, export, and delete.",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "Extras: Layer Architecture Documentation", 1)
    add_paragraph(
        doc,
        "This section is extra documentation for reviewers who want to understand how the system is put together. The main submission details are above.",
    )

    add_heading(doc, "Overall Layering", 2)
    add_paragraph(
        doc,
        "The app is organized around two acquisition pipelines that meet at one shared contact draft. The image path and voice path stay separate because their inputs are different. A photo needs image cleanup, OCR, and line-based parsing. Voice needs audio capture, transcription, spoken-text cleanup, and natural-language parsing.",
    )
    add_paragraph(
        doc,
        "Once either pipeline produces a contact draft, the rest of the application is shared: review form, duplicate check, merge decision, save, relationship linking, group assignment, contact list, VCF export, and graph view.",
    )

    add_heading(doc, "Frontend Layer", 2)
    add_bullets(
        doc,
        [
            "React page layer controls the three main views: Add Contact, Contacts, and Contact Graph.",
            "Input components handle image upload/camera capture, voice recording, and manual field edits.",
            "The review form is intentionally editable because OCR and speech results can be imperfect.",
            "Contact detail components handle export, deletion, group assignment, and manual relationship linking.",
            "The graph view uses a mature graph renderer for pan, zoom, drag, search, labels, and group coloring.",
        ],
    )

    add_heading(doc, "Business Card Pipeline", 2)
    add_numbers(
        doc,
        [
            "The user captures or uploads a business card image.",
            "The React app sends the image to POST /extractions/business-card.",
            "NestJS validates the upload and passes the image through preprocessing variants.",
            "Tesseract runs locally and returns OCR text.",
            "The parser extracts emails, phones, website, name, title, company, address, and relationship hints.",
            "The backend returns a ContactDraft to the frontend.",
            "The user reviews and saves the result.",
        ],
    )

    add_heading(doc, "Voice Pipeline", 2)
    add_numbers(
        doc,
        [
            "The user records audio or enters a spoken-style note.",
            "The frontend sends the audio or transcript to POST /extractions/voice.",
            "The backend uses local speech-to-text for audio input.",
            "The transcript is cleaned so spoken numbers, email phrases, websites, and labels are easier to parse.",
            "The natural-language parser maps the transcript into the same ContactDraft contract.",
            "The user can also fill individual empty fields quickly with the field-level voice button.",
        ],
    )

    add_heading(doc, "Contact Workflow Layer", 2)
    add_paragraph(
        doc,
        "The contact workflow is shared for all input sources. Before saving, duplicate detection checks likely matches by normalized name, email, and phone. If a match exists, the user can keep the old contact, merge new details into it, or create a separate contact.",
    )
    add_paragraph(
        doc,
        "After save or merge, the app can suggest related contacts. Same-family-name contacts and same-company contacts are ranked higher, but the user still chooses whether to link them or skip.",
    )

    add_heading(doc, "Persistence Layer", 2)
    add_paragraph(
        doc,
        "PostgreSQL stores the contact manager data. Drizzle ORM owns the schema definitions, migrations, and database access. Contacts are stored separately from emails, phones, websites, and addresses so one contact can have multiple values for each field.",
    )
    add_paragraph(
        doc,
        "Relationships are stored as contact-to-contact edges. Groups are stored separately from relationships through a many-to-many group membership table. This lets a person belong to multiple groups without forcing every group to be a graph relationship.",
    )

    add_heading(doc, "Graph Layer", 2)
    add_paragraph(
        doc,
        "The graph is contact-only: there are no fake parent nodes for groups. Groups affect color and filtering, while relationships become labeled edges between contacts. This keeps the graph closer to how a user thinks about people.",
    )

    add_heading(doc, "Docker Layer", 2)
    add_paragraph(
        doc,
        "Docker Compose starts PostgreSQL, the NestJS API, the React build served through Nginx, and Adminer. The API startup script applies migrations and seeds only when the database is empty. The web container proxies /api requests to the API container, so the reviewer can use one browser origin for the app.",
    )

    add_heading(doc, "Known Practical Note", 2)
    add_paragraph(
        doc,
        "OCR quality depends heavily on photo quality, card layout, and text contrast. The app treats extracted results as a draft on purpose. The reliable workflow is scan, review, correct if needed, then save.",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
